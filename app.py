import streamlit as st
import os
import time
import uuid
from config import AVAILABLE_MODELS, MODEL_CAPABILITIES, DEFAULT_MODEL
from chat_handler import handle_chat_message
from file_handler import process_uploaded_file
from image_handler import generate_image
from utils import initialize_session_state
from api_utils import get_euron_api_key
from database_handler import DatabaseLogger
import io
from PIL import Image
from fpdf import FPDF
from docx import Document
import base64

def create_pdf(messages):
    """Create a PDF from chat messages."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    # Add a title
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, "ResearchBuddy AI Chat History", ln=True, align='C')
    pdf.ln(10)
    
    # Reset font for content
    pdf.set_font("Arial", size=12)
    
    # Add chat messages
    for message in messages:
        role = message["role"].upper()
        content = message["content"]
        
        # Add role with appropriate styling
        if role == "USER":
            pdf.set_text_color(0, 0, 255)  # Blue for user
        else:
            pdf.set_text_color(0, 128, 0)  # Green for assistant
            
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(200, 10, f"{role}:", ln=True)
        
        # Reset styling for content
        pdf.set_text_color(0, 0, 0)  # Black for content
        pdf.set_font("Arial", size=12)
        
        # Add content with multi-cell for proper wrapping
        pdf.multi_cell(0, 10, content)
        pdf.ln(5)
    
    return pdf.output(dest="S").encode("latin1")

def create_docx(messages):
    """Create a DOCX from chat messages."""
    doc = Document()
    
    # Add a title
    doc.add_heading("ResearchBuddy AI Chat History", 0)
    
    # Add chat messages
    for message in messages:
        role = message["role"].upper()
        content = message["content"]
        
        # Add message with role as heading
        doc.add_heading(f"{role}:", 2)
        doc.add_paragraph(content)
        doc.add_paragraph()  # Add space between messages
    
    # Save to memory
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def get_download_link(file_bytes, file_name, file_type):
    """Generate a download link for the file."""
    b64 = base64.b64encode(file_bytes).decode()
    return f'<a href="data:{file_type};base64,{b64}" download="{file_name}">Download {file_name}</a>'

def initialize_database_logger():
    """Initialize the database logger if not already in session state."""
    if 'db_logger' not in st.session_state:
        st.session_state.db_logger = DatabaseLogger()
    return st.session_state.db_logger

def ensure_session_id():
    """Ensure we have a unique session ID for this chat session."""
    if 'session_id' not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())
    return st.session_state.session_id

def main():
    """Main function to run the Streamlit app."""
    st.set_page_config(page_title="ResearchBuddy AI: A Multi-Model AI Assistant", layout="wide")
    
    st.title("ResearchBuddy AI: A Multi-Model AI Assistant")
    
    # Initialize session state and database
    initialize_session_state()
    db_logger = initialize_database_logger()
    session_id = ensure_session_id()
    
    # Log the session with some browser info if available
    try:
        user_agent = st.experimental_get_query_params().get('user_agent', [None])[0]
        db_logger.log_session(session_id, user_browser=user_agent)
    except:
        db_logger.log_session(session_id)
    
    # Sidebar for model selection and settings
    with st.sidebar:
        st.header("Model Settings")
        
        # Model selection dropdown
        selected_model = st.selectbox(
            "Select AI Model",
            options=list(AVAILABLE_MODELS.keys()),
            index=0,
            key="model_selection"
        )
        
        st.caption(f"Model ID: {AVAILABLE_MODELS[selected_model]}")
        
        # Display model capabilities
        st.subheader("Model Capabilities")
        capabilities = MODEL_CAPABILITIES.get(selected_model, {})
        for capability, supported in capabilities.items():
            st.checkbox(capability, value=supported, disabled=True)
        
        # Check if API key is available
        api_key = get_euron_api_key()
        if not api_key:
            st.warning("API key not found in secrets. Please add it to your .streamlit/secrets.toml file.")
            st.code("""
            # Example secrets.toml file
            [euron]
            api_key = "your-api-key-here"
            """)
        else:
            st.success("Please Select the Model Based on your Desired Task.")
        
        # Temperature slider
        temperature = st.slider("Temperature", min_value=0.0, max_value=1.0, value=0.7, step=0.1)
        st.session_state.temperature = temperature
        
        # Max tokens slider
        max_tokens = st.slider("Max Tokens", min_value=100, max_value=3000, value=1000, step=100)
        st.session_state.max_tokens = max_tokens
        
        # Clear chat button
        if st.button("Clear Chat"):
            st.session_state.messages = []
            st.session_state.uploaded_file_content = None
            st.session_state.uploaded_file_name = None
            st.session_state.uploaded_image = None
            st.experimental_rerun()
        
        # Download chat options
        if st.session_state.messages:
            st.subheader("Download Chat")
            download_format = st.radio("Select Format", ["PDF", "DOCX"])
            
            if st.button("Download Chat History"):
                if download_format == "PDF":
                    pdf_bytes = create_pdf(st.session_state.messages)
                    st.download_button(
                        label="Download as PDF",
                        data=pdf_bytes,
                        file_name="researchbuddy_chat.pdf",
                        mime="application/pdf"
                    )
                else:  # DOCX
                    docx_bytes = create_docx(st.session_state.messages)
                    st.download_button(
                        label="Download as DOCX",
                        data=docx_bytes,
                        file_name="researchbuddy_chat.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        
        # Admin section for database stats (hidden by default)
        with st.expander("Admin Stats", expanded=False):
            if st.button("Refresh Stats"):
                stats = db_logger.get_stats()
                st.write(f"Total Sessions: {stats['total_sessions']}")
                st.write(f"Total Interactions: {stats['total_interactions']}")
                st.write(f"Most Popular Model: {stats['most_popular_model']} ({stats['most_popular_model_count']} uses)")
    
    # Main content area
    col1, col2 = st.columns([3, 1])
    
    with col2:
        # File upload section
        st.subheader("File Upload")
        uploaded_file = st.file_uploader("Upload a file", type=["txt", "pdf", "csv", "xlsx", "jpg", "jpeg", "png"])
        if uploaded_file:
            file_details = process_uploaded_file(uploaded_file)
            st.session_state.uploaded_file_content = file_details["content"]
            st.session_state.uploaded_file_name = file_details["name"]
            
            # If it's an image, store it separately and display it
            if file_details["is_image"]:
                st.session_state.uploaded_image = file_details["image"]
                st.image(file_details["image"], caption=file_details["name"], use_column_width=True)
                st.info("Image uploaded! You can now ask questions about this image.")
            else:
                st.success(f"File uploaded: {file_details['name']}")
        
        # Image generation section
        st.subheader("Image Generation")
        image_prompt = st.text_input("Image Description")
        if st.button("Generate Image") and image_prompt:
            with st.spinner("Generating image..."):
                start_time = time.time()
                image_result = generate_image(
                    image_prompt, 
                    api_key,  # This parameter is now ignored but kept for compatibility
                    selected_model
                )
                execution_time_ms = int((time.time() - start_time) * 1000)
                
                # Log the image generation request
                db_logger.log_interaction(
                    session_id=session_id,
                    model_name=selected_model,
                    model_id=AVAILABLE_MODELS[selected_model],
                    temperature=st.session_state.temperature,
                    max_tokens=st.session_state.max_tokens,
                    user_query=f"[IMAGE GENERATION] {image_prompt}",
                    model_response="Image generated successfully" if "image" in image_result else f"Error: {image_result.get('error')}",
                    has_image=True,
                    execution_time_ms=execution_time_ms
                )
                
                if "image" in image_result:
                    st.image(image_result["image"], caption=image_prompt)
                else:
                    st.error(image_result["error"])
    
    with col1:
        # Chat interface
        st.subheader("Chat")
        
        # Display uploaded file info if any
        if st.session_state.uploaded_file_name:
            st.info(f"Uploaded file: {st.session_state.uploaded_file_name}")
        
        # Display chat messages
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        
        # Input for new message
        user_input = st.chat_input("Ask something...")
        if user_input:
            # Add user message to chat history
            st.session_state.messages.append({"role": "user", "content": user_input})
            
            # Display user message
            with st.chat_message("user"):
                st.markdown(user_input)
            
            # Generate AI response
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    message_placeholder = st.empty()
                    
                    # Check if an image is uploaded and the selected model supports image analysis
                    has_image = st.session_state.uploaded_image is not None
                    model_supports_images = MODEL_CAPABILITIES.get(selected_model, {}).get("Image Analysis", False)
                    
                    # If image is uploaded and model doesn't support images, add a warning
                    if has_image and not model_supports_images:
                        st.warning(f"Note: {selected_model} doesn't fully support image analysis. For best results with images, try using Google Gemini 2.5 Pro Exp.")
                    
                    # Time the execution
                    start_time = time.time()
                    
                    response = handle_chat_message(
                        user_input,
                        st.session_state.messages,
                        selected_model,
                        AVAILABLE_MODELS[selected_model],
                        api_key,  # This parameter is now ignored but kept for compatibility
                        st.session_state.temperature,
                        st.session_state.max_tokens,
                        st.session_state.uploaded_file_content,
                        st.session_state.uploaded_image if has_image else None
                    )
                    
                    # Calculate execution time
                    execution_time_ms = int((time.time() - start_time) * 1000)
                    
                    # Log the interaction to database
                    db_logger.log_interaction(
                        session_id=session_id,
                        model_name=selected_model,
                        model_id=AVAILABLE_MODELS[selected_model],
                        temperature=st.session_state.temperature,
                        max_tokens=st.session_state.max_tokens,
                        user_query=user_input,
                        model_response=response,
                        has_file=st.session_state.uploaded_file_name is not None,
                        file_name=st.session_state.uploaded_file_name,
                        has_image=has_image,
                        execution_time_ms=execution_time_ms
                    )
                    
                    message_placeholder.markdown(response)
            
            # Add assistant response to chat history
            st.session_state.messages.append({"role": "assistant", "content": response})
            
            # Download latest response option
            if st.button("Download Latest Response"):
                # Create a document with just the latest response
                last_message = st.session_state.messages[-1]
                
                # Let user choose format
                download_col1, download_col2 = st.columns(2)
                with download_col1:
                    if st.button("Download as PDF"):
                        pdf_bytes = create_pdf([last_message])
                        st.download_button(
                            label="Download PDF",
                            data=pdf_bytes,
                            file_name="researchbuddy_response.pdf",
                            mime="application/pdf"
                        )
                with download_col2:
                    if st.button("Download as DOCX"):
                        docx_bytes = create_docx([last_message])
                        st.download_button(
                            label="Download DOCX",
                            data=docx_bytes,
                            file_name="researchbuddy_response.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

if __name__ == "__main__":
    main()
